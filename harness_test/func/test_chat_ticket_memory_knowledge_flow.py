from __future__ import annotations

from pathlib import Path

from docx import Document

from harness_test.fixture.app_client import auth_headers


def test_chat_suggests_ticket_without_creating_one(api_client, monkeypatch):
    """验证处理对话`suggests`工单`without``creating``one`。

    :param api_client: 隔离测试环境提供的 FastAPI 测试客户端与辅助数据；类型由调用方及当前处理场景决定。
    :param monkeypatch: pytest 提供的运行时替换与环境变量修改夹具；类型由调用方及当前处理场景决定。
    :return: 无返回值；函数通过副作用、断言或异常完成其职责。
    """
    import agent_server.graph_flow.graph_nodes as graph_nodes

    monkeypatch.setattr(
        graph_nodes,
        "decide_with_llm",
        lambda question, context: {"answer": "please create a ticket", "needs_ticket": True, "title": "network fault"},
    )

    headers = auth_headers(api_client, "alice")

    chat = api_client.post("/api/chat", json={"message": "my router is broken"}, headers=headers)
    tickets = api_client.get("/api/tickets", headers=headers)
    history = api_client.get("/api/chat/history", headers=headers)

    assert chat.status_code == 200, chat.text
    data = chat.json()["data"]
    assert data["ticket_id"] is None
    assert data["ticket_suggestion"] == {
        "recommended": True,
        "title": "network fault",
        "content": "my router is broken",
        "answer": "please create a ticket",
    }
    assert tickets.json()["data"]["items"] == []
    assert history.json()["data"]["items"][0]["ticket_id"] is None


def test_user_creates_pending_ticket_and_admin_approves(api_client):
    """验证用户`creates``pending`工单`and`管理员`approves`。

    :param api_client: 隔离测试环境提供的 FastAPI 测试客户端与辅助数据；类型由调用方及当前处理场景决定。
    :return: 无返回值；函数通过副作用、断言或异常完成其职责。
    """
    employee_headers = auth_headers(api_client, "alice")
    admin_headers = auth_headers(api_client, "root", role="admin")

    created = api_client.post(
        "/api/tickets",
        json={"title": "router fault", "content": "router is broken", "answer": "please check hardware"},
        headers=employee_headers,
    )

    assert created.status_code == 200, created.text
    ticket = created.json()["data"]
    assert ticket["status"] == "pending"
    assert ticket["metadata"]

    employee_patch = api_client.patch(f"/api/tickets/{ticket['id']}", json={"status": "approved"}, headers=employee_headers)
    admin_patch = api_client.patch(f"/api/tickets/{ticket['id']}", json={"status": "approved"}, headers=admin_headers)

    assert employee_patch.status_code == 403
    assert admin_patch.status_code == 200, admin_patch.text
    assert admin_patch.json()["data"]["status"] == "approved"


def test_chat_uses_recent_history_as_context(api_client, monkeypatch):
    """验证处理对话`uses``recent`历史记录`as``context`。

    :param api_client: 隔离测试环境提供的 FastAPI 测试客户端与辅助数据；类型由调用方及当前处理场景决定。
    :param monkeypatch: pytest 提供的运行时替换与环境变量修改夹具；类型由调用方及当前处理场景决定。
    :return: 无返回值；函数通过副作用、断言或异常完成其职责。
    """
    import agent_server.core.db as db
    import agent_server.graph_flow.graph_nodes as graph_nodes

    captured: dict[str, str] = {}

    def fake_decide(question: str, context: str):
        """`fake`生成决策。

        :param question: 函数处理所需的“问题”数据，类型为 ``str``。
        :param context: 函数处理所需的“`context`”数据，类型为 ``str``。
        :return: 返回`fake`生成决策得到的处理结果；具体类型由实际执行分支决定。
        """
        captured["context"] = context
        return {"answer": "memory answer", "needs_ticket": False, "title": "memory"}

    monkeypatch.setattr(graph_nodes, "decide_with_llm", fake_decide)
    headers = auth_headers(api_client, "alice")
    user = api_client.get("/api/auth/me", headers=headers).json()["data"]
    db.create_chat_history(user["id"], "first question", "first answer")

    response = api_client.post("/api/chat", json={"message": "what did I ask before"}, headers=headers)

    assert response.status_code == 200, response.text
    assert "first question" in captured["context"]
    assert "first answer" in captured["context"]


def test_admin_deletes_uploaded_document(api_client, tmp_path, monkeypatch):
    """验证管理员`deletes``uploaded`文档。

    :param api_client: 隔离测试环境提供的 FastAPI 测试客户端与辅助数据；类型由调用方及当前处理场景决定。
    :param tmp_path: pytest 为当前测试提供的隔离临时目录；类型由调用方及当前处理场景决定。
    :param monkeypatch: pytest 提供的运行时替换与环境变量修改夹具；类型由调用方及当前处理场景决定。
    :return: 无返回值；函数通过副作用、断言或异常完成其职责。
    """
    import agent_server.core.db as db
    import agent_server.tools.business_tools as business_tools

    calls: list[str] = []
    monkeypatch.setattr(business_tools, "rebuild_index", lambda: (None, [], None))

    admin_headers = auth_headers(api_client, "root", role="admin")
    source = tmp_path / "datas" / "delete-me.docx"
    source.write_bytes(b"fake docx")
    doc = db.upsert_doc(str(source), "delete-me.docx", chunk_count=3)

    response = api_client.delete(f"/api/knowledge/{doc['id']}", headers=admin_headers)

    assert response.status_code == 200, response.text
    assert response.json()["data"]["deleted"]["id"] == doc["id"]
    assert not source.exists()
    assert db.list_docs() == []


def test_admin_reads_original_docx_text(api_client, tmp_path):
    """管理员可以读取已入库 Word 文档的正文。

    :param api_client: 隔离测试客户端。
    :param tmp_path: pytest 临时目录。
    :return: 无返回值；断言接口返回原始正文。
    """
    import agent_server.core.db as db

    admin_headers = auth_headers(api_client, "root", role="admin")
    source = tmp_path / "datas" / "employee-handbook.docx"
    document = Document()
    document.add_paragraph("员工手册原始正文")
    document.add_paragraph("第二段内容")
    document.save(source)
    doc = db.upsert_doc(str(source), source.name, chunk_count=2)

    response = api_client.get(f"/api/knowledge/{doc['id']}/content", headers=admin_headers)

    assert response.status_code == 200, response.text
    assert response.json()["data"] == {
        "id": doc["id"],
        "title": source.name,
        "content": "员工手册原始正文\n\n第二段内容",
    }
