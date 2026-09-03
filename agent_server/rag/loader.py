from __future__ import annotations

import hashlib
import logging
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass, field
from functools import lru_cache
from io import BytesIO
from pathlib import Path
from typing import Any, Iterable

import fitz
import pdfplumber
from docx import Document

from common.constants import DATAS_DIR, PROJECT_ROOT, VLM_MODEL_DIR


LOGGER = logging.getLogger(__name__)
SUPPORTED_SUFFIXES = {".pdf", ".docx", ".doc"}
VLM_REQUIRED_FILES = ("config.json",)


@dataclass(frozen=True)
class DocumentBlock:
    source_path: str
    text: str
    block_type: str = "text"
    page: int | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True)
class LoadedDocument:
    source_path: str
    blocks: list[DocumentBlock]
    checksum: str
    table_blocks: int = 0
    image_blocks: int = 0


_OCR_ENGINE: Any | None = None
_OCR_FAILED = False


def scan_source_files(source_dir: str | Path = DATAS_DIR) -> list[Path]:
    """扫描源文件`files`。

    :param source_dir: 函数处理所需的“源文件`dir`”数据，类型为 ``str | Path``。
    :return: 返回扫描源文件`files`得到的结果，返回类型为 ``list[Path]``。
    """
    root = Path(source_dir)
    if not root.exists():
        return []

    files: list[Path] = []
    for path in root.rglob("*"):
        if not path.is_file():
            continue
        relative_parts = {part.lower() for part in path.relative_to(root).parts}
        if "chroma" in relative_parts or path.name.lower() == "app.db":
            continue
        if path.suffix.lower() in SUPPORTED_SUFFIXES:
            files.append(path)
    return sorted(files, key=lambda item: str(item).lower())


def load_documents(source_dir: str | Path = DATAS_DIR, enable_ocr: bool = True) -> list[LoadedDocument]:
    """加载`documents`。

    :param source_dir: 函数处理所需的“源文件`dir`”数据，类型为 ``str | Path``。
    :param enable_ocr: 函数处理所需的“`enable`OCR”数据，类型为 ``bool``。
    :return: 返回加载`documents`得到的结果，返回类型为 ``list[LoadedDocument]``。
    """
    loaded: list[LoadedDocument] = []
    for path in scan_source_files(source_dir):
        try:
            if path.suffix.lower() == ".pdf":
                loaded.append(load_pdf(path, enable_ocr=enable_ocr))
            elif path.suffix.lower() == ".docx":
                loaded.append(load_docx(path, enable_ocr=enable_ocr))
            elif path.suffix.lower() == ".doc":
                converted = convert_doc_to_docx(path)
                if converted is None:
                    LOGGER.warning("Skip old .doc file because conversion failed: %s", path)
                    continue
                loaded.append(load_docx(converted, enable_ocr=enable_ocr, original_path=path))
        except Exception as exc:
            LOGGER.warning("Skip document after extraction failure: %s (%s)", path, exc)
    return loaded


def load_pdf(path: str | Path, enable_ocr: bool = True) -> LoadedDocument:
    """加载`pdf`。

    :param path: 目标文件或目录路径，类型为 ``str | Path``。
    :param enable_ocr: 函数处理所需的“`enable`OCR”数据，类型为 ``bool``。
    :return: 返回加载`pdf`得到的结果，返回类型为 ``LoadedDocument``。
    """
    source = Path(path)
    blocks: list[DocumentBlock] = []

    with pdfplumber.open(source) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                blocks.append(_block(source, text, "text", page_index))
            for table in page.extract_tables() or []:
                markdown = table_to_markdown(table)
                if markdown:
                    blocks.append(_block(source, markdown, "table", page_index))

    if enable_ocr:
        blocks.extend(_extract_pdf_image_blocks(source))

    return _loaded(source, blocks)


def load_docx(path: str | Path, enable_ocr: bool = True, original_path: Path | None = None) -> LoadedDocument:
    """加载`docx`。

    :param path: 目标文件或目录路径，类型为 ``str | Path``。
    :param enable_ocr: 函数处理所需的“`enable`OCR”数据，类型为 ``bool``。
    :param original_path: 函数处理所需的“`original`路径”数据，类型为 ``Path | None``。
    :return: 返回加载`docx`得到的结果，返回类型为 ``LoadedDocument``。
    """
    source = Path(original_path or path)
    doc = Document(path)
    blocks: list[DocumentBlock] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(_block(source, text, "text"))

    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        markdown = table_to_markdown(rows)
        if markdown:
            blocks.append(_block(source, markdown, "table"))

    if enable_ocr:
        blocks.extend(_extract_docx_image_blocks(Path(path), source))

    return _loaded(source, blocks)


def table_to_markdown(rows: Iterable[Iterable[Any]]) -> str:
    """转换表格`to``markdown`。

    :param rows: 需要写入、转换或聚合的多行数据，类型为 ``Iterable[Iterable[Any]]``。
    :return: 返回转换表格`to``markdown`得到的结果，返回类型为 ``str``。
    """
    cleaned = [[" ".join(str(cell or "").split()) for cell in row] for row in rows]
    cleaned = [row for row in cleaned if any(row)]
    if not cleaned:
        return ""

    width = max(len(row) for row in cleaned)
    padded = [row + [""] * (width - len(row)) for row in cleaned]
    header = padded[0]
    separator = ["---"] * width
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    for row in padded[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def convert_doc_to_docx(path: Path) -> Path | None:
    """转换知识文档`to``docx`。

    :param path: 目标文件或目录路径，类型为 ``Path``。
    :return: 返回转换知识文档`to``docx`得到的结果，返回类型为 ``Path | None``。
    """
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice is None:
        LOGGER.warning("Cannot convert .doc because soffice/libreoffice is not on PATH: %s", path)
        return None

    output_dir = Path(tempfile.mkdtemp(prefix="doc-convert-"))
    result = subprocess.run(
        [soffice, "--headless", "--convert-to", "docx", "--outdir", str(output_dir), str(path)],
        capture_output=True,
        text=True,
        check=False,
        timeout=120,
    )
    converted = output_dir / f"{path.stem}.docx"
    if result.returncode != 0 or not converted.exists():
        LOGGER.warning("soffice conversion failed for %s: %s", path, result.stderr.strip())
        return None
    return converted


def _extract_pdf_image_blocks(path: Path) -> list[DocumentBlock]:
    """提取`pdf`图片`blocks`。

    :param path: 目标文件或目录路径，类型为 ``Path``。
    :return: 返回提取`pdf`图片`blocks`得到的结果，返回类型为 ``list[DocumentBlock]``。
    """
    blocks: list[DocumentBlock] = []
    with fitz.open(path) as doc:
        for page_index in range(len(doc)):
            page = doc[page_index]
            for image_index, image in enumerate(page.get_images(full=True), start=1):
                xref = image[0]
                image_data = doc.extract_image(xref)
                payload = image_data.get("image")
                ext = image_data.get("ext", "png")
                if not payload:
                    continue
                text = ocr_image_bytes(payload, ext)
                if text:
                    blocks.append(
                        _block(
                            path,
                            f"[插图内容: {text}]",
                            "image",
                            page_index + 1,
                            {"image_index": image_index, "ocr": True},
                        )
                    )
    return blocks


def _extract_docx_image_blocks(path: Path, source_path: Path) -> list[DocumentBlock]:
    """提取`docx`图片`blocks`。

    :param path: 目标文件或目录路径，类型为 ``Path``。
    :param source_path: 函数处理所需的“源文件路径”数据，类型为 ``Path``。
    :return: 返回提取`docx`图片`blocks`得到的结果，返回类型为 ``list[DocumentBlock]``。
    """
    blocks: list[DocumentBlock] = []
    with zipfile.ZipFile(path) as archive:
        media_names = [name for name in archive.namelist() if name.startswith("word/media/")]
        for image_index, name in enumerate(media_names, start=1):
            payload = archive.read(name)
            ext = Path(name).suffix.lstrip(".") or "png"
            text = ocr_image_bytes(payload, ext)
            if text:
                blocks.append(
                    _block(
                        source_path,
                        f"[插图内容: {text}]",
                        "image",
                        None,
                        {"image_index": image_index, "media_name": name, "ocr": True},
                    )
                )
    return blocks


def ocr_image_bytes(payload: bytes, ext: str = "png") -> str:
    """OCR图片`bytes`。

    :param payload: 函数处理所需的“`payload`”数据，类型为 ``bytes``。
    :param ext: 函数处理所需的“`ext`”数据，类型为 ``str``。
    :return: 返回OCR图片`bytes`得到的结果，返回类型为 ``str``。
    """
    suffix = "." + ext.lower().lstrip(".")
    if suffix == ".emf":
        vector_text = extract_vector_image_text(payload)
        if vector_text:
            return enhance_image_text(payload, ext, vector_text)

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(payload)
        tmp_path = Path(tmp.name)
    try:
        return enhance_image_text(payload, ext, ocr_image_file(tmp_path))
    finally:
        tmp_path.unlink(missing_ok=True)


def vlm_available() -> bool:
    """视觉语言模型`available`。

    :return: 返回视觉语言模型`available`得到的结果，返回类型为 ``bool``。
    """
    if not VLM_MODEL_DIR.exists():
        return False
    has_required = all((VLM_MODEL_DIR / name).exists() for name in VLM_REQUIRED_FILES)
    has_weight = any(
        (VLM_MODEL_DIR / name).exists()
        for name in ("model.safetensors", "pytorch_model.bin", "model.safetensors.index.json")
    )
    has_processor = any(
        (VLM_MODEL_DIR / name).exists()
        for name in ("preprocessor_config.json", "processor_config.json", "tokenizer_config.json")
    )
    return has_required and has_weight and has_processor


def vlm_enabled() -> bool:
    """视觉语言模型`enabled`。

    :return: 返回视觉语言模型`enabled`得到的结果，返回类型为 ``bool``。
    """
    value = os.getenv("VLM_ENABLED", "true").strip().lower()
    return value in {"1", "true", "yes", "on"} and vlm_available() and _module_available("bitsandbytes")


def vlm_status() -> dict[str, object]:
    """视觉语言模型获取状态。

    :return: 返回视觉语言模型获取状态得到的结果，返回类型为 ``dict[str, object]``。
    """
    available = vlm_available()
    bitsandbytes_available = _module_available("bitsandbytes")
    skip_reason = None
    if not available:
        skip_reason = f"未发现完整本地 VLM 模型，插图语义增强继续使用 PaddleOCR 文本: {VLM_MODEL_DIR}"
    elif not bitsandbytes_available:
        skip_reason = "本地 VLM 模型已就位，但当前环境未安装 bitsandbytes，load_in_4bit=True 的真实加载暂不能运行。"
    return {
        "model_path": str(VLM_MODEL_DIR),
        "available": available,
        "enabled": vlm_enabled(),
        "load_in_4bit": True,
        "bitsandbytes_available": bitsandbytes_available,
        "skip_reason": skip_reason,
        "todo": None if available and bitsandbytes_available else "安装 4-bit 依赖后再运行 VLM 插图语义验收。",
    }


def build_vlm_quantization_config() -> Any:
    """构建视觉语言模型`quantization`配置。

    :return: 返回构建视觉语言模型`quantization`配置得到的结果，返回类型为 ``Any``。
    """
    try:
        from transformers import BitsAndBytesConfig

        return BitsAndBytesConfig(load_in_4bit=True)
    except Exception:
        return {"load_in_4bit": True}


def enhance_image_text(payload: bytes, ext: str = "png", ocr_text: str = "") -> str:
    """增强图片文本。

    :param payload: 函数处理所需的“`payload`”数据，类型为 ``bytes``。
    :param ext: 函数处理所需的“`ext`”数据，类型为 ``str``。
    :param ocr_text: 函数处理所需的“OCR文本”数据，类型为 ``str``。
    :return: 返回增强图片文本得到的结果，返回类型为 ``str``。
    """
    caption = caption_image_with_vlm(payload, ext) if vlm_enabled() else ""
    parts = [part for part in (ocr_text.strip(), caption.strip()) if part]
    return " / ".join(_dedupe(parts))


def caption_image_with_vlm(payload: bytes, ext: str = "png") -> str:
    """生成说明文本图片`with`视觉语言模型。

    :param payload: 函数处理所需的“`payload`”数据，类型为 ``bytes``。
    :param ext: 函数处理所需的“`ext`”数据，类型为 ``str``。
    :return: 返回生成说明文本图片`with`视觉语言模型得到的结果，返回类型为 ``str``。
    """
    if not vlm_enabled():
        return ""
    try:
        model, processor = get_vlm_model_and_processor()
        from PIL import Image

        image = Image.open(BytesIO(payload)).convert("RGB")
        prompt = "请用简洁中文描述这张企业制度插图中的流程、节点和关键文字。"
        messages = [{"role": "user", "content": [{"type": "image"}, {"type": "text", "text": prompt}]}]
        text = _apply_vlm_chat_template(processor, messages)
        inputs = processor(text=[text], images=[image], return_tensors="pt")
        if hasattr(model, "device"):
            inputs = {key: value.to(model.device) if hasattr(value, "to") else value for key, value in inputs.items()}
        output_ids = model.generate(**inputs, max_new_tokens=128)
        prompt_length = inputs["input_ids"].shape[-1]
        generated = output_ids[:, prompt_length:]
        captions = processor.batch_decode(generated, skip_special_tokens=True, clean_up_tokenization_spaces=False)
        return str(captions[0]).strip() if captions else ""
    except Exception as exc:
        LOGGER.warning("VLM semantic caption failed; using OCR text only: %s", exc)
        return ""


def _apply_vlm_chat_template(processor: Any, messages: list[dict[str, Any]]) -> str:
    """应用视觉语言模型处理对话`template`。

    :param processor: 函数处理所需的“`processor`”数据，类型为 ``Any``。
    :param messages: 函数处理所需的“`messages`”数据，类型为 ``list[dict[str, Any]]``。
    :return: 返回应用视觉语言模型处理对话`template`得到的结果，返回类型为 ``str``。
    """
    template_kwargs = {"tokenize": False, "add_generation_prompt": True}
    if hasattr(processor, "apply_chat_template"):
        try:
            return processor.apply_chat_template(messages, **template_kwargs)
        except Exception as exc:
            if "chat template" not in str(exc).lower():
                raise

    tokenizer = getattr(processor, "tokenizer", None)
    if tokenizer is not None and hasattr(tokenizer, "apply_chat_template"):
        return tokenizer.apply_chat_template(messages, **template_kwargs)

    prompt = messages[0]["content"][1]["text"]
    return (
        "<|im_start|>user\n"
        f"<|vision_start|><|image_pad|><|vision_end|>{prompt}<|im_end|>\n"
        "<|im_start|>assistant\n"
    )


@lru_cache(maxsize=1)
def get_vlm_model_and_processor() -> tuple[Any, Any]:
    """获取视觉语言模型模型`and``processor`。

    :return: 返回获取视觉语言模型模型`and``processor`得到的结果，返回类型为 ``tuple[Any, Any]``。
    :raises FileNotFoundError: 当代码中对应的校验或操作失败条件成立时抛出。
    """
    if not vlm_available():
        raise FileNotFoundError(f"Local VLM model not found or incomplete: {VLM_MODEL_DIR}")
    from transformers import AutoProcessor, Qwen2_5_VLForConditionalGeneration

    model = Qwen2_5_VLForConditionalGeneration.from_pretrained(
        str(VLM_MODEL_DIR),
        local_files_only=True,
        trust_remote_code=True,
        device_map="auto",
        quantization_config=build_vlm_quantization_config(),
    )
    processor = AutoProcessor.from_pretrained(
        str(VLM_MODEL_DIR),
        local_files_only=True,
        trust_remote_code=True,
    )
    return model, processor


def ocr_image_file(path: str | Path) -> str:
    """OCR图片文件。

    :param path: 目标文件或目录路径，类型为 ``str | Path``。
    :return: 返回OCR图片文件得到的结果，返回类型为 ``str``。
    """
    global _OCR_FAILED
    if _OCR_FAILED:
        return ""
    try:
        engine = get_ocr_engine()
        result = engine.ocr(str(path))
        tokens = _collect_ocr_text(result)
        return " / ".join(tokens)
    except Exception as exc:
        _OCR_FAILED = True
        LOGGER.warning("PaddleOCR failed; image OCR will be skipped for this run: %s", exc)
        return ""


def get_ocr_engine() -> Any:
    """获取OCR`engine`。

    :return: 返回获取OCR`engine`得到的结果，返回类型为 ``Any``。
    """
    global _OCR_ENGINE
    if _OCR_ENGINE is None:
        cache_dir = PROJECT_ROOT / ".cache" / "paddle"
        cache_dir.mkdir(parents=True, exist_ok=True)
        os.environ.setdefault("PADDLE_HOME", str(cache_dir))
        os.environ.setdefault("PADDLE_PDX_CACHE_HOME", str(cache_dir / "paddlex"))
        os.environ.setdefault("XDG_CACHE_HOME", str(PROJECT_ROOT / ".cache"))

        from paddleocr import PaddleOCR

        kwargs = {
            "lang": "ch",
            "use_doc_orientation_classify": False,
            "use_doc_unwarping": False,
            "use_textline_orientation": False,
        }
        _OCR_ENGINE = PaddleOCR(**kwargs)
    return _OCR_ENGINE


def extract_vector_image_text(payload: bytes) -> str:
    """提取向量图片文本。

    :param payload: 函数处理所需的“`payload`”数据，类型为 ``bytes``。
    :return: 返回提取向量图片文本得到的结果，返回类型为 ``str``。
    """
    texts: list[str] = []
    for encoding in ("utf-16le", "gbk", "utf-8"):
        decoded = payload.decode(encoding, errors="ignore")
        matches = re.findall(r"[\u4e00-\u9fffA-Za-z0-9（）()、，。；：:/\\\-]{2,}", decoded)
        texts.extend(item.strip() for item in matches if _meaningful_vector_text(item))
    return " / ".join(_dedupe(texts))


def _meaningful_vector_text(text: str) -> bool:
    """`meaningful`向量文本。

    :param text: 需要校验、解析或转换的文本，类型为 ``str``。
    :return: 返回`meaningful`向量文本得到的结果，返回类型为 ``bool``。
    """
    if len(text.strip()) < 2:
        return False
    chinese_count = sum(1 for char in text if "\u4e00" <= char <= "\u9fff")
    return chinese_count >= 2


def _collect_ocr_text(node: Any) -> list[str]:
    """收集OCR文本。

    :param node: 待分析的 Python 函数语法树节点，类型为 ``Any``。
    :return: 返回收集OCR文本得到的结果，返回类型为 ``list[str]``。
    """
    texts: list[str] = []
    if node is None:
        return texts
    if isinstance(node, dict):
        for key in ("rec_texts", "text", "transcription"):
            value = node.get(key)
            if isinstance(value, str) and value.strip():
                texts.append(value.strip())
            elif isinstance(value, list):
                texts.extend(str(item).strip() for item in value if str(item).strip())
        for value in node.values():
            if isinstance(value, (list, tuple, dict)):
                texts.extend(_collect_ocr_text(value))
        return _dedupe(texts)
    if isinstance(node, (list, tuple)):
        if len(node) >= 2 and isinstance(node[1], (list, tuple)) and node[1]:
            candidate = node[1][0]
            if isinstance(candidate, str) and candidate.strip():
                texts.append(candidate.strip())
        for item in node:
            if isinstance(item, (list, tuple, dict)):
                texts.extend(_collect_ocr_text(item))
        return _dedupe(texts)
    return texts


def _block(
    source: Path,
    text: str,
    block_type: str,
    page: int | None = None,
    metadata: dict[str, Any] | None = None,
) -> DocumentBlock:
    """内容块。

    :param source: 函数处理所需的“源文件”数据，类型为 ``Path``。
    :param text: 需要校验、解析或转换的文本，类型为 ``str``。
    :param block_type: 函数处理所需的“内容块`type`”数据，类型为 ``str``。
    :param page: 函数处理所需的“页面”数据，类型为 ``int | None``。
    :param metadata: 函数处理所需的“元数据”数据，类型为 ``dict[str, Any] | None``。
    :return: 返回内容块得到的结果，返回类型为 ``DocumentBlock``。
    """
    return DocumentBlock(
        source_path=str(source),
        page=page,
        block_type=block_type,
        text=text.strip(),
        metadata=metadata or {},
    )


def _loaded(source: Path, blocks: list[DocumentBlock]) -> LoadedDocument:
    """`loaded`。

    :param source: 函数处理所需的“源文件”数据，类型为 ``Path``。
    :param blocks: 函数处理所需的“`blocks`”数据，类型为 ``list[DocumentBlock]``。
    :return: 返回`loaded`得到的结果，返回类型为 ``LoadedDocument``。
    """
    return LoadedDocument(
        source_path=str(source),
        blocks=blocks,
        checksum=_sha256(source),
        table_blocks=sum(1 for block in blocks if block.block_type == "table"),
        image_blocks=sum(1 for block in blocks if block.block_type == "image"),
    )


def _sha256(path: Path) -> str:
    """`sha256`。

    :param path: 目标文件或目录路径，类型为 ``Path``。
    :return: 返回`sha256`得到的结果，返回类型为 ``str``。
    """
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _dedupe(items: Iterable[str]) -> list[str]:
    """`dedupe`。

    :param items: 需要批量处理的数据项，类型为 ``Iterable[str]``。
    :return: 返回`dedupe`得到的结果，返回类型为 ``list[str]``。
    """
    seen: set[str] = set()
    result: list[str] = []
    for item in items:
        text = item.strip()
        if text and text not in seen:
            seen.add(text)
            result.append(text)
    return result


def _module_available(module_name: str) -> bool:
    """`module``available`。

    :param module_name: 函数处理所需的“`module``name`”数据，类型为 ``str``。
    :return: 返回`module``available`得到的结果，返回类型为 ``bool``。
    """
    try:
        __import__(module_name)
        return True
    except Exception:
        return False
