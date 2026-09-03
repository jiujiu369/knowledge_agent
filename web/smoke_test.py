from __future__ import annotations

import json
import os
import sys
import time
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Any

import requests
from docx import Document

from web.frontend_api import auth_headers, parse_sse_events


BASE_URL = os.getenv("KNOWLEDGE_AGENT_API_BASE_URL", "http://127.0.0.1:8000").rstrip("/")
PASSWORD = "Passw0rd!"
ADMIN_USERNAME_ENV = "KNOWLEDGE_AGENT_SMOKE_ADMIN_USERNAME"
ADMIN_PASSWORD_ENV = "KNOWLEDGE_AGENT_SMOKE_ADMIN_PASSWORD"


def _json(response: requests.Response) -> dict[str, Any]:
    """`json`。

    :param response: 需要解析或转换的 HTTP 响应对象，类型为 ``requests.Response``。
    :return: 返回`json`得到的结果，返回类型为 ``dict[str, Any]``。
    :raises AssertionError: 当代码中对应的校验或操作失败条件成立时抛出。
    """
    try:
        payload = response.json()
    except ValueError as exc:
        raise AssertionError(f"HTTP {response.status_code} returned non-json: {response.text[:200]}") from exc
    assert payload.get("code") == "ok", payload
    return payload


def _register(username: str, role: str) -> None:
    """注册。

    :param username: 用于定位账户的用户名，类型为 ``str``。
    :param role: 用于权限判断的用户角色标识，类型为 ``str``。
    :return: 无返回值；函数通过副作用、断言或异常完成其职责。
    """
    response = requests.post(
        f"{BASE_URL}/api/auth/register",
        json={"username": username, "password": PASSWORD, "role": role},
        timeout=20,
    )
    assert response.status_code in {200, 400}, response.text
    if response.status_code == 400:
        assert "already exists" in response.text, response.text


def _login(username: str, password: str = PASSWORD) -> str:
    """执行登录。

    :param username: 用于定位账户的用户名，类型为 ``str``。
    :param password: 用于登录验证的账户密码，类型为 ``str``。
    :return: 返回执行登录得到的结果，返回类型为 ``str``。
    """
    response = requests.post(
        f"{BASE_URL}/api/auth/login",
        json={"username": username, "password": password},
        timeout=20,
    )
    assert response.status_code == 200, response.text
    data = _json(response)["data"]
    assert isinstance(data.get("token"), str) and data["token"], data
    assert data.get("role") in {"employee", "admin", "hr", "finance", "ops"}, data
    return data["token"]


def _make_docx() -> Path:
    """创建`docx`。

    :return: 返回创建`docx`得到的结果，返回类型为 ``Path``。
    """
    document = Document()
    document.add_heading("M3 前端上传冒烟文档", level=1)
    document.add_paragraph("这是一份用于 Streamlit 上传页冒烟测试的知识库文档。")
    document.add_paragraph("测试内容：员工可查询制度，管理员可上传文档并重建索引。")
    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        path = Path(tmp.name)
    document.save(path)
    return path


def _prepare_admin(_legacy_admin: str | None = None) -> tuple[str, dict[str, str]]:
    """准备管理员。

    :param _legacy_admin: 兼容旧调用方的参数，不参与管理员身份来源，类型为 ``str | None``。
    :return: 返回准备管理员得到的结果，返回类型为 ``tuple[str, dict[str, str]]``。
    """
    admin = os.getenv(ADMIN_USERNAME_ENV)
    admin_password = os.getenv(ADMIN_PASSWORD_ENV)
    assert admin and admin_password, (
        f"Web smoke requires pre-created admin credentials in {ADMIN_USERNAME_ENV} and {ADMIN_PASSWORD_ENV}"
    )
    token = _login(admin, admin_password)
    return token, auth_headers(token)


def main() -> None:
    """执行当前模块的主流程并协调各项处理步骤。

    :return: 无返回值；函数通过副作用、断言或异常完成其职责。
    """
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    suffix = str(int(time.time()))
    employee = f"m3_emp_{suffix}"

    print("M3 smoke: health")
    response = requests.get(f"{BASE_URL}/health", timeout=10)
    assert response.status_code == 200, response.text
    assert response.json()["status"] == "ok", response.text

    print("M3 smoke: login pre-created admin")
    admin_token, admin_headers = _prepare_admin()

    print("M3 smoke: rebuild knowledge index")
    response = requests.post(f"{BASE_URL}/api/knowledge/rebuild", headers=admin_headers, timeout=300)
    assert response.status_code == 200, response.text
    rebuild_data = _json(response)["data"]
    assert isinstance(rebuild_data.get("stats"), list), rebuild_data
    print(f"knowledge_stats={len(rebuild_data['stats'])}")

    print("M3 smoke: register/login employee")
    _register(employee, "employee")
    employee_token = _login(employee)
    employee_headers = auth_headers(employee_token)

    print("M3 smoke: chat SSE")
    with requests.post(
        f"{BASE_URL}/api/chat/stream",
        json={"message": "差旅报销上限多少"},
        headers=employee_headers,
        stream=True,
        timeout=(10, 240),
    ) as response:
        assert response.status_code == 200, response.text
        events = list(parse_sse_events(response.iter_lines(decode_unicode=True)))
    event_names = [item["event"] for item in events]
    assert "tool" in event_names, events
    assert "done" in event_names, events
    done = next(item["data"] for item in events if item["event"] == "done")
    assert isinstance(done.get("answer"), str) and done["answer"], done
    assert "tool_events" in done and isinstance(done["tool_events"], list), done
    print(f"chat_events={json.dumps(event_names, ensure_ascii=False)}")
    print(f"ticket_id={done.get('ticket_id')}")

    print("M3 smoke: query ticket list")
    response = requests.get(f"{BASE_URL}/api/tickets", headers=employee_headers, timeout=20)
    assert response.status_code == 200, response.text
    tickets = _json(response)["data"]
    assert isinstance(tickets.get("items"), list), tickets
    print(f"ticket_count={len(tickets['items'])}")

    print("M3 smoke: admin export ticket stat")
    response = requests.post(f"{BASE_URL}/api/tools/export_ticket_stat", json={}, headers=admin_headers, timeout=20)
    assert response.status_code == 200, response.text
    stat = _json(response)["data"]
    assert "total" in stat and "by_status" in stat, stat
    print(f"export_total={stat['total']}")

    print("M3 smoke: admin upload knowledge file")
    docx_path = _make_docx()
    try:
        with docx_path.open("rb") as handle:
            response = requests.post(
                f"{BASE_URL}/api/knowledge/upload",
                files={"file": ("m3_streamlit_smoke.docx", handle)},
                headers=admin_headers,
                timeout=120,
            )
        assert response.status_code == 200, response.text
        upload_data = _json(response)["data"]
        assert upload_data.get("source_path", "").endswith("m3_streamlit_smoke.docx"), upload_data
        print("upload=ok")
    finally:
        docx_path.unlink(missing_ok=True)

    print("M3 自检通过")


if __name__ == "__main__":
    main()
